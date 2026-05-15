"""
格式转换技能。
支持 Word、PDF、TXT 等常见文档格式互转。
Author: 破执
Date: 2026-05-15
"""

import os
import subprocess
import shutil
from pathlib import Path

from logger import get_logger

logger = get_logger('format_converter')

# 技能元数据
SKILL_NAME = "format_converter"
SKILL_DESCRIPTION = "文档格式转换，支持 Word、PDF、TXT 等格式互转。"
SKILL_TRIGGER = "当需要将文档从一种格式转换为另一种格式时使用，如 Word 转 PDF、PDF 转 TXT 等。"
SKILL_CATEGORY = "io"
SKILL_REQUIRES_CONFIRMATION = False
SKILL_PARAMETERS = [
    {
        "name": "input_path",
        "type": "string",
        "description": "输入文件路径"
    },
    {
        "name": "output_path",
        "type": "string",
        "description": "输出文件路径（可选，默认与输入文件同目录）",
        "default": ""
    },
    {
        "name": "output_format",
        "type": "string",
        "description": "目标格式: pdf, docx, txt, html, md",
        "default": ""
    }
]


def _get_format_from_path(path: str) -> str:
    """从路径获取文件格式。"""
    return Path(path).suffix.lower().lstrip(".")


def _validate_paths(input_path: str, output_path: str, output_format: str) -> tuple:
    """
    验证路径参数。
    返回 (是否有效, input_abs, output_abs, target_format 或 错误信息)
    """
    if not input_path or not input_path.strip():
        return False, None, None, "输入文件路径不能为空"

    input_abs = os.path.abspath(input_path.strip())
    if not os.path.exists(input_abs):
        return False, None, None, f"输入文件不存在: {input_abs}"

    input_format = _get_format_from_path(input_abs)

    # 确定目标格式
    target_format = output_format.strip().lower() if output_format else ""
    if not target_format and output_path:
        target_format = _get_format_from_path(output_path.strip())

    if not target_format:
        return False, None, None, "无法确定目标格式，请提供 output_format 或带后缀的 output_path"

    # 确定输出路径
    if output_path and output_path.strip():
        output_abs = os.path.abspath(output_path.strip())
    else:
        base = os.path.splitext(input_abs)[0]
        output_abs = f"{base}.{target_format}"

    # 创建输出目录
    output_dir = os.path.dirname(output_abs)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    return True, input_abs, output_abs, target_format


def _convert_with_pandoc(input_path: str, output_path: str, target_format: str) -> str:
    """使用 pandoc 进行转换。"""
    pandoc_path = shutil.which("pandoc")
    if not pandoc_path:
        return "❌ 未找到 pandoc。请安装: https://pandoc.org/installing.html"

    # 格式映射
    format_map = {
        "docx": "docx",
        "pdf": "pdf",
        "txt": "plain",
        "html": "html",
        "md": "markdown",
        "epub": "epub",
        "rtf": "rtf",
    }

    pandoc_fmt = format_map.get(target_format, target_format)

    try:
        cmd = [pandoc_path, input_path, "-o", output_path, "-t", pandoc_fmt]
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60,
            encoding="utf-8",
            errors="replace"
        )

        if result.returncode != 0:
            return f"❌ pandoc 转换失败: {result.stderr}"

        return f"✅ 转换成功: {os.path.basename(input_path)} → {os.path.basename(output_path)}"

    except subprocess.TimeoutExpired:
        return "❌ 转换超时 (超过 60s)"
    except Exception as e:
        return f"❌ 转换异常: {str(e)}"


def _convert_txt_to_docx(input_path: str, output_path: str) -> str:
    """TXT 转 DOCX（使用 python-docx）。"""
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        with open(input_path, "r", encoding="utf-8", errors="replace") as f:
            for line in f:
                line = line.rstrip("\n\r")
                if line.strip():
                    p = doc.add_paragraph(line)
                else:
                    doc.add_paragraph()

        doc.save(output_path)
        return f"✅ 转换成功: {os.path.basename(input_path)} → {os.path.basename(output_path)}"

    except ImportError:
        return "❌ 未安装 python-docx。请运行: pip install python-docx"
    except Exception as e:
        return f"❌ 转换失败: {str(e)}"


def _convert_docx_to_txt(input_path: str, output_path: str) -> str:
    """DOCX 转 TXT（使用 python-docx）。"""
    try:
        from docx import Document

        doc = Document(input_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(paragraphs))

        return f"✅ 转换成功: {os.path.basename(input_path)} → {os.path.basename(output_path)}"

    except ImportError:
        return "❌ 未安装 python-docx。请运行: pip install python-docx"
    except Exception as e:
        return f"❌ 转换失败: {str(e)}"


def _convert_pdf_to_txt(input_path: str, output_path: str) -> str:
    """PDF 转 TXT（使用 PyPDF2 或 pdfplumber）。"""
    # 先尝试 pdfplumber（效果更好）
    try:
        import pdfplumber

        text_parts = []
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n\n".join(text_parts))

        return f"✅ 转换成功: {os.path.basename(input_path)} → {os.path.basename(output_path)}"

    except ImportError:
        pass

    # 降级到 PyPDF2
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(input_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n\n".join(text_parts))

        return f"✅ 转换成功: {os.path.basename(input_path)} → {os.path.basename(output_path)}"

    except ImportError:
        return "❌ 未安装 PDF 处理库。请运行: pip install pdfplumber 或 pip install PyPDF2"
    except Exception as e:
        return f"❌ 转换失败: {str(e)}"


def _convert_txt_to_pdf(input_path: str, output_path: str) -> str:
    """TXT 转 PDF（使用 reportlab）。"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # 尝试注册中文字体
        try:
            pdfmetrics.registerFont(TTFont("SimSun", "simsun.ttc"))
            font_name = "SimSun"
        except:
            try:
                pdfmetrics.registerFont(TTFont("SimSun", "SimSun.ttf"))
                font_name = "SimSun"
            except:
                font_name = "Helvetica"

        c = canvas.Canvas(output_path, pagesize=A4)
        width, height = A4
        margin = 50
        line_height = 14
        y = height - margin

        with open(input_path, "r", encoding="utf-8", errors="replace") as f:
            for line in f:
                line = line.rstrip("\n\r")
                if y < margin + line_height:
                    c.showPage()
                    y = height - margin

                c.setFont(font_name, 10)
                # 处理长行换行
                text_width = width - 2 * margin
                words = []
                current_line = ""
                for char in line:
                    test_line = current_line + char
                    if c.stringWidth(test_line, font_name, 10) > text_width:
                        words.append(current_line)
                        current_line = char
                    else:
                        current_line = test_line
                if current_line:
                    words.append(current_line)

                for word in words:
                    if y < margin + line_height:
                        c.showPage()
                        y = height - margin
                        c.setFont(font_name, 10)
                    c.drawString(margin, y, word)
                    y -= line_height

                y -= line_height  # 段落间距

        c.save()
        return f"✅ 转换成功: {os.path.basename(input_path)} → {os.path.basename(output_path)}"

    except ImportError:
        return "❌ 未安装 reportlab。请运行: pip install reportlab"
    except Exception as e:
        return f"❌ 转换失败: {str(e)}"


def execute(input_path: str, output_path: str = "", output_format: str = "", **kwargs) -> str:
    """
    执行文档格式转换。

    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径（可选）
        output_format: 目标格式（可选）
        **kwargs: 额外参数（忽略）

    Returns:
        转换结果
    """
    valid, input_abs, output_abs, target = _validate_paths(input_path, output_path, output_format)
    if not valid:
        return f"❌ {target}"

    input_format = _get_format_from_path(input_abs)
    logger.info(f"格式转换: {input_abs} ({input_format}) → {output_abs} ({target})")

    # 检查是否需要转换
    if input_format == target:
        return f"⚠️ 输入格式与目标格式相同 ({target})，无需转换"

    # 根据格式组合选择转换方法
    combo = (input_format, target)

    # 优先使用专用转换器
    if combo == ("txt", "docx"):
        return _convert_txt_to_docx(input_abs, output_abs)
    elif combo == ("docx", "txt"):
        return _convert_docx_to_txt(input_abs, output_abs)
    elif combo == ("pdf", "txt"):
        return _convert_pdf_to_txt(input_abs, output_abs)
    elif combo == ("txt", "pdf"):
        return _convert_txt_to_pdf(input_abs, output_abs)

    # 其他格式使用 pandoc
    return _convert_with_pandoc(input_abs, output_abs, target)
